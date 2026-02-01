import asyncio
import os
import requests
import json
import pandas as pd
from pypdf import PdfReader
from docx import Document
import config
from browser_agent import BrowserAgent
from extractor import JobExtractor

# Filenames
# Filenames
CV_FILE = config.CV_FILENAME
TEMPLATE_FILE = config.TEMPLATE_FILENAME
JOBS_FILE = os.path.join(config.DATA_DIR, "jobs_found.xlsx")
LLM_API_URL = f"{config.LLM_API_BASE}/chat/completions"

async def fetch_full_job_description(url):
    """Fetches and cleans the full job description from the URL."""
    print(f"  [Browser] Fetching full JD from: {url}")
    agent = BrowserAgent()
    extractor = JobExtractor()
    
    try:
        await agent.start()
        await agent.navigate_to(url)
        # Wait a bit for dynamic content
        await agent.human_delay(2, 4)
        html = await agent.get_page_content()
        cleaned_text = extractor.clean_html(html)
        return cleaned_text
    except Exception as e:
        print(f"  [Browser] Error fetching JD: {e}")
        return ""
    finally:
        await agent.stop()

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
        return ""

def generate_cover_letter_body(cv_text, job_description, company_name, job_position):
    """Uses LLM to write the body of the cover letter."""
    
    system_prompt = (
        "You are a professional career coach. "
        "Write the BODY of a cover letter. "
        "STRICTLY DO NOT include the header, date, 'Dear Hiring Manager', 'Subject', or 'Sincerely'. "
        "The template ALREADY has these. If you include them, it will look duplicated. "
        "Start directly with the opening sentence (e.g., 'I am writing to express my strong interest...'). "
        "Write only the 3-4 paragraphs of the body content. "
        "Tailor the content to match the candidate's CV skills with the Job Description requirements. "
        "Keep it professional and persuasive. Limit to 300 words."
    )
    
    user_prompt = f"""
    CANDIDATE CV:
    {cv_text[:3000]}
    
    TARGET JOB:
    Position: {job_position}
    Company: {company_name}
    Description: {job_description}
    
    Write the body paragraphs now.
    """
    
    payload = {
        "model": config.LLM_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.7
    }
    
    try:
        print("  [LLM] Generating cover letter content...")
        response = requests.post(LLM_API_URL, headers={"Content-Type": "application/json"}, json=payload, timeout=120)
        
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            print(f"  [LLM] API Error: {response.text}")
            return "Error generating content."
            
    except Exception as e:
        print(f"  [LLM] Connection Error: {e}")
        return "Error connecting to LLM."

def create_cover_letter_doc(body_text, output_filename, company, position):
    """Creates a new Word doc based on the template by replacing placeholders."""
    from datetime import datetime
    
    try:
        if os.path.exists(TEMPLATE_FILE):
             doc = Document(TEMPLATE_FILE)
             print(f"  [Doc] Loaded template: {TEMPLATE_FILE}")
        else:
             print("Please ensure the template file exists.")
             return False

        current_date = datetime.now().strftime("%d %B %Y")
        
        # Clean body text: Remove \r, ensure paragraphs are separated by \n\n
        clean_body = body_text.replace("\r", "").strip()
        body_paragraphs = [p.strip() for p in clean_body.split('\n') if p.strip()]

        # Replacements for simple single-line placeholders
        replacements = {
            "<Today Date>": current_date,
            "<Company Name>": company,
            "<Position>": position
        }

        # 1. Handle standard replacements
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

        # 2. Handle Body Paragraphs
        # We look for <Paragraph 1>, insert the real body paragraphs before it, and then remove <Paragraph 1/2/3>
        # Note: Modifying list while iterating is risky. We collect actions first or restart index.
        # Simpler: Iterate and when we find <Paragraph 1>, do the insertion.
        
        target_para = None
        for paragraph in doc.paragraphs:
            if "<Paragraph 1>" in paragraph.text:
                target_para = paragraph
                break
        
        if target_para:
            print(f"  [Debug] Raw LLM Body (First 200 chars): {body_text[:200]}...")
            
            # Filter body paragraphs to remove any accidentally generated salutations
            final_paragraphs = []
            for i, p in enumerate(body_paragraphs):
                p_lower = p.lower()
                
                # Check for Salutation at start
                if i == 0:
                    start_phrases = ["dear hiring manager", "dear", "to the hiring manager", "to whom it may concern"]
                    for phrase in start_phrases:
                        if p_lower.startswith(phrase):
                            # Remove the phrase and any following punctuation/whitespace
                            # Naive strip: find the existing case-insensitive match length
                            # Better: just split and keep the rest?
                            # Or simple: if it's short (just the salutation), skip it.
                            # If it's long (Salutation + Body), split it.
                            if len(p) < 50: 
                                p = "" # It's just a salutation line
                            else:
                                # Start of body on same line? Try to strip the prefix.
                                # Finding index of match is tricky case-insensitively, 
                                # but usually LLM puts it on separate line.
                                # Let's assume if it's long, we keep it but warn? 
                                # actually, let's just skip "Dear..." check for long paragraphs to be safe
                                # unless it strictly starts with it.
                                pass 
                            break
                
                # Check for Sign-off at end
                if i >= len(body_paragraphs) - 2:
                    sign_offs = ["sincerely", "regards", "best regards", "yours truly", "best"]
                    for phrase in sign_offs:
                        if p_lower.startswith(phrase):
                             p = ""
                             break
                
                if p.strip():
                    final_paragraphs.append(p)

            print(f"  [Debug] Final Paragraphs count: {len(final_paragraphs)}")

            # Insert generated paragraphs before the target
            # And explicitly ensure they are NOT bold.
            for para_text in final_paragraphs:
                new_p = target_para.insert_paragraph_before(para_text)
                # Clear direct formatting to rely on style defaults (usually Normal)
                # And explicitly set bold to False to override any inheritance from the placeholder check
                for run in new_p.runs:
                    run.bold = False
                
                # Force style to 'Normal' if possible, or just reset bold
                new_p.style = doc.styles['Normal']
            
            # Now extract specific paragraphs to delete (the placeholders)
            # Deleting paragraphs in python-docx is not straight forward (need to remove xml element)
            def delete_paragraph(paragraph):
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

            # Find and delete all body placeholders
            # Need to iterate copy of paragraphs list to avoid issues
            for paragraph in list(doc.paragraphs):
                if any(ph in paragraph.text for ph in ["<Paragraph 1>", "<Paragraph 2>", "<Paragraph 3>"]):
                    delete_paragraph(paragraph)

        doc.save(output_filename)
        print(f"  [Doc] Saved cover letter to: {output_filename}")
        return True
        
    except Exception as e:
        print(f"Error creating Word doc: {e}")
        return False

async def main():
    if not os.path.exists(CV_FILE):
        print(f"Error: CV file '{CV_FILE}' not found.")
        return
    
    if not os.path.exists(JOBS_FILE):
        print(f"Error: Jobs file '{JOBS_FILE}' not found. Run main.py first.")
        return

    # 1. Read CV
    print("Reading CV...")
    cv_text = extract_text_from_pdf(CV_FILE)
    if not cv_text:
        return

    # 2. Read Jobs
    print("Reading Job List...")
    df = pd.read_excel(JOBS_FILE)
    if df.empty:
        print("No jobs found in Excel.")
        return

    # 3. Process All Jobs
    print(f"Found {len(df)} jobs. Starting generation...")
    
    # Create output directory
    output_dir = "Generated_Cover_Letters"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for index, row in df.iterrows():
        company = row.get('Company', 'Unknown Company')
        position = row.get('Job Position', 'Unknown Position')
        link = row.get('Link to post', '')
        
        print(f"\n--- Processing Job {index + 1}/{len(df)}: {position} at {company} ---")
        
        description = ""
        if link and str(link).startswith('http'):
            # Only fetch if we haven't effectively cached it or if we want fresh
            # For now, simplistic approach: always fetch
            description = await fetch_full_job_description(link)
        else:
            print("  [Warning] No valid link found, using excel summary.")
            description = row.get('Full Job Description', '')

        if not description:
            print("  [Skip] Could not get job description. Skipping.")
            continue

        # 4. Generate Content
        body_text = generate_cover_letter_body(cv_text, description, company, position)
        if not body_text or "Error" in body_text:
             print("  [Skip] Failed to generate body text.")
             continue

        # 5. Create Document
        safe_company = "".join(c for c in str(company) if c.isalnum() or c in (' ', '_', '-')).strip()
        safe_position = "".join(c for c in str(position) if c.isalnum() or c in (' ', '_', '-')).strip()
        filename = f"{safe_company}_{safe_position}.docx".replace(" ", "_")
        output_path = os.path.join(output_dir, filename)
        
        create_cover_letter_doc(body_text, output_path, company, position)

if __name__ == "__main__":
    asyncio.run(main())
